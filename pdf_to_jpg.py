import os
from pdf2image import convert_from_path
from PIL import Image, ImageDraw, ImageFont
import docx

# Define function to save the first page of a PDF as an image
def save_first_page_pdf_as_image(pdf_path, output_dir):
    images = convert_from_path(pdf_path, first_page=1, last_page=1)  # Convert only the first page
    for i, image in enumerate(images):
        img_output_path = os.path.join(output_dir, f"{os.path.basename(pdf_path).replace('.pdf', '')}.jpg")
        image.save(img_output_path, 'JPEG')
        print(f'Saved: {img_output_path}')

# Define function to save the first page of a DOCX as an image
def save_first_page_docx_as_image(docx_path, output_path):
    doc = docx.Document(docx_path)
    # Create a blank image
    img = Image.new('RGB', (800, 600), color='white')
    draw = ImageDraw.Draw(img)

    # Set a font (default font)
    font = ImageFont.load_default()

    y_text = 10
    if doc.paragraphs:  # Only draw if there are paragraphs
        para = doc.paragraphs[0]  # Get the first paragraph
        draw.text((10, y_text), para.text, font=font, fill=(0, 0, 0))

    img.save(output_path.replace('.docx', '.jpg'))
    print(f'Saved: {output_path.replace(".docx", ".jpg")}')

# Directory paths
input_dir = 'D:/NCKH/Loaded_Images/pdf'  # Input directory with PDF and DOCX files
output_dir = 'D:/NCKH/Loaded_Images/jpg'  # Output directory for JPG images

# Create output directory if it doesn't exist
os.makedirs(output_dir, exist_ok=True)

# Convert all PDF and DOCX files in the input directory to JPG
for filename in os.listdir(input_dir):
    file_path = os.path.join(input_dir, filename)
    if filename.endswith('.pdf'):
        save_first_page_pdf_as_image(file_path, output_dir)
    elif filename.endswith('.docx'):
        save_first_page_docx_as_image(file_path, os.path.join(output_dir, filename))
